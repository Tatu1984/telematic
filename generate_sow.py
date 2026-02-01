#!/usr/bin/env python3
"""
Generate Statement of Work (SoW) document for FleetTrack Pro Telematics Platform
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle, Circle, Polygon
import numpy as np

# Output directory
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def create_data_flow_diagram():
    """Create data flow diagram with proper spacing"""
    fig, ax = plt.subplots(1, 1, figsize=(18, 14))
    ax.set_xlim(0, 18)
    ax.set_ylim(-1, 14)
    ax.axis('off')
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('#FAFAFA')

    colors = {
        'user': '#4A90D9',
        'frontend': '#50C878',
        'api': '#FF8C42',
        'auth': '#9B59B6',
        'database': '#E74C3C',
        'external': '#7F8C8D',
        'arrow': '#2C3E50'
    }

    # Title
    ax.text(9, 13.3, 'FleetTrack Pro - Data Flow Diagram', fontsize=20, fontweight='bold',
            ha='center', va='center', color='#2C3E50')

    def draw_box(x, y, w, h, color, label, sublabel=None):
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05,rounding_size=0.2",
                             facecolor=color, edgecolor='#2C3E50', linewidth=2, alpha=0.9)
        ax.add_patch(box)
        if sublabel:
            ax.text(x + w/2, y + h/2 + 0.2, label, fontsize=11, fontweight='bold',
                    ha='center', va='center', color='white')
            ax.text(x + w/2, y + h/2 - 0.2, sublabel, fontsize=8,
                    ha='center', va='center', color='white', alpha=0.9)
        else:
            ax.text(x + w/2, y + h/2, label, fontsize=11, fontweight='bold',
                    ha='center', va='center', color='white')

    def draw_arrow(start, end, label=None, color='#2C3E50', curve=0):
        style = "Simple,tail_width=0.5,head_width=4,head_length=6"
        connectionstyle = f"arc3,rad={curve}" if curve != 0 else "arc3,rad=0"
        arrow = FancyArrowPatch(start, end, arrowstyle=style, color=color,
                                connectionstyle=connectionstyle, linewidth=1.5, alpha=0.7)
        ax.add_patch(arrow)
        if label:
            mid_x = (start[0] + end[0]) / 2 + (curve * 2)
            mid_y = (start[1] + end[1]) / 2 + (0.35 if curve == 0 else curve * 1.5)
            ax.text(mid_x, mid_y, label, fontsize=8, ha='center', va='center',
                    color='#2C3E50', style='italic',
                    bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.9, edgecolor='#BDC3C7'))

    # Layer labels
    ax.text(0.3, 11.8, 'USERS', fontsize=9, fontweight='bold', color='#7F8C8D', ha='left')
    ax.text(0.3, 8.8, 'FRONTEND', fontsize=9, fontweight='bold', color='#7F8C8D', ha='left')
    ax.text(0.3, 5.8, 'API LAYER', fontsize=9, fontweight='bold', color='#7F8C8D', ha='left')
    ax.text(0.3, 3.0, 'SERVICES', fontsize=9, fontweight='bold', color='#7F8C8D', ha='left')
    ax.text(0.3, 0.3, 'DATABASE', fontsize=9, fontweight='bold', color='#7F8C8D', ha='left')

    # Layer 1: Users (y=10.5 to 12)
    draw_box(1.5, 10.5, 3, 1.3, colors['user'], 'Fleet Manager', 'Web Browser')
    draw_box(5.5, 10.5, 3, 1.3, colors['user'], 'Driver', 'Mobile/Web')
    draw_box(9.5, 10.5, 3, 1.3, colors['user'], 'Admin', 'Web Portal')
    draw_box(14, 10.5, 3, 1.3, colors['external'], 'IoT Devices', 'GPS/ELD/OBD')

    # Layer 2: Frontend (y=7.5 to 9)
    draw_box(3, 7.5, 8, 1.5, colors['frontend'], 'Next.js Frontend', 'React 19 + Tailwind CSS + Leaflet Maps + Recharts')

    # Layer 3: API & Auth (y=4.5 to 6)
    draw_box(1, 4.5, 3.5, 1.5, colors['auth'], 'NextAuth', 'JWT Sessions')
    draw_box(5.5, 4.5, 5, 1.5, colors['api'], 'REST API Layer', 'Rate Limiting + Zod Validation')
    draw_box(11.5, 4.5, 4, 1.5, colors['api'], 'API Routes', '30+ Endpoints')

    # Layer 4: Services (y=1.8 to 3.3)
    draw_box(0.5, 1.8, 2.8, 1.5, colors['auth'], 'Auth Service', 'RBAC + Audit')
    draw_box(3.7, 1.8, 2.8, 1.5, colors['api'], 'Fleet Service', 'Vehicles/Drivers')
    draw_box(6.9, 1.8, 2.8, 1.5, colors['api'], 'Tracking', 'Real-time GPS')
    draw_box(10.1, 1.8, 2.8, 1.5, colors['api'], 'Analytics', 'Metrics/Reports')
    draw_box(13.3, 1.8, 2.8, 1.5, colors['api'], 'Compliance', 'ELD/DOT Logs')

    # Database (y=-0.5 to 0.8)
    draw_box(4.5, -0.5, 9, 1.5, colors['database'], 'PostgreSQL Database', 'Prisma ORM - 24 Tables - Indexed for Performance')

    # Arrows: Users to Frontend
    draw_arrow((3, 10.5), (5, 9), 'HTTP/HTTPS')
    draw_arrow((7, 10.5), (7, 9), 'HTTP/HTTPS')
    draw_arrow((11, 10.5), (9, 9), 'HTTP/HTTPS')

    # IoT to API
    draw_arrow((15.5, 10.5), (13.5, 6), 'Telemetry', curve=-0.2)

    # Frontend to API/Auth
    draw_arrow((5, 7.5), (2.75, 6), 'Auth')
    draw_arrow((7, 7.5), (8, 6), 'API Calls')

    # Auth to Service
    draw_arrow((2.75, 4.5), (1.9, 3.3), 'Validate')

    # API to Services
    draw_arrow((6.5, 4.5), (5.1, 3.3), 'CRUD')
    draw_arrow((8.5, 4.5), (8.3, 3.3), 'Query')
    draw_arrow((11.5, 4.8), (11.5, 3.3), 'Aggregate')
    draw_arrow((13.5, 4.5), (14.7, 3.3), 'Logs')

    # Services to Database
    draw_arrow((1.9, 1.8), (6, 1), 'R/W', curve=0.15)
    draw_arrow((5.1, 1.8), (7.5, 1), 'CRUD')
    draw_arrow((8.3, 1.8), (9, 1), 'Query')
    draw_arrow((11.5, 1.8), (10.5, 1), 'Agg', curve=-0.1)
    draw_arrow((14.7, 1.8), (12, 1), 'Logs', curve=-0.15)

    # Legend
    legend_y = 12.5
    legend_x = 14.5
    ax.add_patch(Rectangle((legend_x, legend_y - 2.2), 3.3, 2.5, facecolor='white', edgecolor='#BDC3C7', alpha=0.9, linewidth=1))
    ax.text(legend_x + 1.65, legend_y + 0.1, 'Legend', fontsize=9, fontweight='bold', ha='center', color='#2C3E50')

    legend_items = [
        (colors['user'], 'Users/Clients'),
        (colors['frontend'], 'Frontend'),
        (colors['api'], 'API/Services'),
        (colors['auth'], 'Authentication'),
        (colors['database'], 'Database'),
        (colors['external'], 'External'),
    ]

    for i, (color, label) in enumerate(legend_items):
        row = i // 2
        col = i % 2
        x = legend_x + 0.15 + col * 1.6
        y = legend_y - 0.5 - row * 0.55
        ax.add_patch(Rectangle((x, y), 0.35, 0.35, facecolor=color, edgecolor='#2C3E50', linewidth=1))
        ax.text(x + 0.45, y + 0.17, label, fontsize=7, va='center', color='#2C3E50')

    plt.tight_layout()
    plt.savefig(os.path.join(OUTPUT_DIR, 'data_flow_diagram.png'), dpi=150, bbox_inches='tight',
                facecolor='#FAFAFA', edgecolor='none', pad_inches=0.3)
    plt.close()
    print("Created: data_flow_diagram.png")

def create_architecture_diagram():
    """Create software architecture diagram with no overlapping"""
    fig, ax = plt.subplots(1, 1, figsize=(20, 16))
    ax.set_xlim(0, 20)
    ax.set_ylim(-1, 16)
    ax.axis('off')
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('#FAFAFA')

    colors = {
        'presentation': '#3498DB',
        'application': '#2ECC71',
        'business': '#F39C12',
        'data': '#E74C3C',
        'infrastructure': '#9B59B6',
        'external': '#7F8C8D',
    }

    # Title
    ax.text(10, 15.3, 'FleetTrack Pro - Software Architecture', fontsize=22, fontweight='bold',
            ha='center', va='center', color='#2C3E50')
    ax.text(10, 14.7, 'Multi-Tenant SaaS Fleet Management Platform', fontsize=13,
            ha='center', va='center', color='#7F8C8D', style='italic')

    def draw_layer(y, height, color, label):
        box = FancyBboxPatch((0.5, y), 19, height, boxstyle="round,pad=0.02,rounding_size=0.1",
                             facecolor=color, edgecolor='#2C3E50', linewidth=2, alpha=0.12)
        ax.add_patch(box)
        ax.text(0.8, y + height - 0.35, label, fontsize=11, fontweight='bold',
                ha='left', va='top', color='#2C3E50')

    def draw_component(x, y, w, h, color, label, items=None):
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.15",
                             facecolor=color, edgecolor='#2C3E50', linewidth=1.5, alpha=0.9)
        ax.add_patch(box)

        if items:
            ax.text(x + w/2, y + h - 0.3, label, fontsize=10, fontweight='bold',
                    ha='center', va='top', color='white')
            for i, item in enumerate(items):
                ax.text(x + w/2, y + h - 0.6 - (i * 0.28), f"- {item}", fontsize=7,
                        ha='center', va='top', color='white', alpha=0.95)
        else:
            ax.text(x + w/2, y + h/2, label, fontsize=10, fontweight='bold',
                    ha='center', va='center', color='white')

    # Layer 1: Presentation (y=12 to 14)
    draw_layer(12, 2.2, colors['presentation'], 'PRESENTATION LAYER')
    draw_component(1.2, 12.3, 2.8, 1.6, colors['presentation'], 'Dashboard', ['Stats Cards', 'Quick Actions'])
    draw_component(4.3, 12.3, 2.8, 1.6, colors['presentation'], 'Live Tracking', ['Leaflet Maps', 'Real-time'])
    draw_component(7.4, 12.3, 2.8, 1.6, colors['presentation'], 'Fleet Mgmt', ['Vehicles', 'Drivers'])
    draw_component(10.5, 12.3, 2.8, 1.6, colors['presentation'], 'ELD Compliance', ['HOS Logs', 'Certification'])
    draw_component(13.6, 12.3, 2.8, 1.6, colors['presentation'], 'Analytics', ['Charts', 'Reports'])
    draw_component(16.7, 12.3, 2.5, 1.6, colors['presentation'], 'Admin Panel', ['Orgs', 'Users'])

    # Layer 2: Application (y=9 to 11.5)
    draw_layer(9, 2.7, colors['application'], 'APPLICATION LAYER (Next.js 16)')
    draw_component(1.2, 9.35, 3.5, 2, colors['application'], 'React Components', ['UI Library', 'Zustand State', 'React Query'])
    draw_component(5, 9.35, 3.5, 2, colors['application'], 'Page Router', ['App Router', 'Middleware', 'Layouts'])
    draw_component(8.8, 9.35, 3.5, 2, colors['application'], 'API Routes', ['REST Endpoints', 'Rate Limiting', 'Zod Validation'])
    draw_component(12.6, 9.35, 3.5, 2, colors['application'], 'Auth (NextAuth)', ['JWT Sessions', 'RBAC', 'Credentials'])
    draw_component(16.4, 9.35, 3, 2, colors['application'], 'Error Handling', ['Sentry', 'Pino Logging'])

    # Layer 3: Business Logic (y=6 to 8.5)
    draw_layer(6, 2.7, colors['business'], 'BUSINESS LOGIC LAYER')
    draw_component(1.2, 6.35, 2.6, 2, colors['business'], 'Fleet Mgmt', ['Vehicle CRUD', 'Driver CRUD'])
    draw_component(4.1, 6.35, 2.6, 2, colors['business'], 'Tracking', ['GPS Processing', 'Telemetry'])
    draw_component(7, 6.35, 2.6, 2, colors['business'], 'Trip Service', ['Route Planning', 'Trip Logging'])
    draw_component(9.9, 6.35, 2.6, 2, colors['business'], 'ELD Service', ['HOS Tracking', 'Certification'])
    draw_component(12.8, 6.35, 2.6, 2, colors['business'], 'Safety', ['Incidents', 'Alerts'])
    draw_component(15.7, 6.35, 2.1, 2, colors['business'], 'Analytics', ['Aggregation'])
    draw_component(18.1, 6.35, 1.3, 2, colors['business'], 'Geo', ['Zones'])

    # Layer 4: Data Access (y=3.2 to 5.5)
    draw_layer(3.2, 2.5, colors['data'], 'DATA ACCESS LAYER')
    draw_component(1.2, 3.55, 4, 1.8, colors['data'], 'Prisma ORM', ['Type-safe Queries', 'Migrations'])
    draw_component(5.5, 3.55, 4, 1.8, colors['data'], 'Repository Pattern', ['CRUD Operations', 'Transactions'])
    draw_component(9.8, 3.55, 4, 1.8, colors['data'], 'Query Builders', ['Filters', 'Pagination', 'Sorting'])
    draw_component(14.1, 3.55, 5.3, 1.8, colors['data'], 'Data Models (24 Tables)', ['Orgs, Users, Vehicles', 'Drivers, Trips, ELD...'])

    # Layer 5: Infrastructure (y=0.3 to 2.8)
    draw_layer(0.3, 2.6, colors['infrastructure'], 'INFRASTRUCTURE LAYER')
    draw_component(1.2, 0.65, 3.2, 1.9, colors['infrastructure'], 'PostgreSQL', ['Primary Database', 'Performance Indexes'])
    draw_component(4.7, 0.65, 3.2, 1.9, colors['infrastructure'], 'Vercel', ['Serverless Deploy', 'Edge Network'])
    draw_component(8.2, 0.65, 3.2, 1.9, colors['infrastructure'], 'Security', ['Rate Limiting', 'CORS Headers'])
    draw_component(11.7, 0.65, 3.2, 1.9, colors['infrastructure'], 'Monitoring', ['Audit Logs', 'Sentry Tracking'])
    draw_component(15.2, 0.65, 2.3, 1.9, colors['external'], 'IoT', ['GPS Devices', 'ELD Units'])
    draw_component(17.8, 0.65, 1.6, 1.9, colors['external'], 'Maps', ['Leaflet'])

    # Vertical connection lines between layers
    def draw_connection(x, y1, y2):
        ax.plot([x, x], [y1, y2], color='#95A5A6', linewidth=1.5, alpha=0.4, linestyle='--')

    # Connections Presentation -> Application
    for x in [2.6, 5.7, 8.8, 11.9, 15, 17.95]:
        draw_connection(x, 12.3, 11.5)

    # Connections Application -> Business
    for x in [2.95, 6.75, 10.55, 14.35, 17.9]:
        draw_connection(x, 9.35, 8.7)

    # Connections Business -> Data
    for x in [2.5, 5.4, 8.3, 11.2, 14.1, 16.75, 18.75]:
        draw_connection(x, 6.35, 5.7)

    # Connections Data -> Infrastructure
    for x in [3.2, 7.5, 11.8, 16.75]:
        draw_connection(x, 3.55, 2.9)

    # Legend at bottom
    ax.text(1, -0.6, 'Architecture Layers:', fontsize=10, fontweight='bold', color='#2C3E50')
    legend_items = [
        (colors['presentation'], 'Presentation'),
        (colors['application'], 'Application'),
        (colors['business'], 'Business Logic'),
        (colors['data'], 'Data Access'),
        (colors['infrastructure'], 'Infrastructure'),
        (colors['external'], 'External Services')
    ]

    for i, (color, label) in enumerate(legend_items):
        x = 5 + i * 2.5
        ax.add_patch(Rectangle((x, -0.75), 0.4, 0.4, facecolor=color, edgecolor='#2C3E50', alpha=0.9, linewidth=1))
        ax.text(x + 0.5, -0.55, label, fontsize=9, va='center', color='#2C3E50')

    plt.tight_layout()
    plt.savefig(os.path.join(OUTPUT_DIR, 'architecture_diagram.png'), dpi=150, bbox_inches='tight',
                facecolor='#FAFAFA', edgecolor='none', pad_inches=0.3)
    plt.close()
    print("Created: architecture_diagram.png")

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level, number=None):
    """Add a numbered heading"""
    if number:
        heading = doc.add_heading(f"{number}. {text}", level)
    else:
        heading = doc.add_heading(text, level)
    return heading

def create_sow_document():
    """Create the Statement of Work document"""
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============ TITLE PAGE ============
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title_run = title.add_run("STATEMENT OF WORK")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("FleetTrack Pro")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(36)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 153)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run("Enterprise Fleet Management & Telematics Platform")
    tagline_run.font.size = Pt(16)
    tagline_run.font.color.rgb = RGBColor(102, 102, 102)
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Project info box
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("Project Type:", "SaaS Web Application"),
        ("Industry:", "Fleet Management & Telematics"),
        ("Platform:", "Web-Based (Desktop & Mobile Responsive)"),
        ("Document Version:", "1.0"),
        ("Date:", "January 2026")
    ]

    for i, (label, value) in enumerate(info_data):
        cell1 = info_table.rows[i].cells[0]
        cell2 = info_table.rows[i].cells[1]
        cell1.text = label
        cell2.text = value
        cell1.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ============ TABLE OF CONTENTS ============
    toc_heading = doc.add_heading("Table of Contents", 1)

    toc_items = [
        ("1.", "Executive Summary", 3),
        ("2.", "Feature List", 3),
        ("   2.1", "Core Fleet Management", 4),
        ("   2.2", "ELD Compliance", 4),
        ("   2.3", "Safety & Analytics", 4),
        ("   2.4", "Administration", 4),
        ("3.", "Technology Stack", 5),
        ("4.", "Data Flow Diagram", 6),
        ("5.", "Software Architecture", 7),
        ("6.", "Plan of Action (Phase-wise Development)", 8),
        ("7.", "Required Team Setup", 10),
        ("8.", "Project Timeline", 11),
        ("9.", "API Documentation", 12),
        ("10.", "Database Schema", 15),
        ("11.", "Appendix", 18),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"{num}  {title}").bold = num.strip().endswith('.')
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        p.add_run(f"\t{page}")

    doc.add_page_break()

    # ============ 1. EXECUTIVE SUMMARY ============
    doc.add_heading("1. Executive Summary", 1)

    exec_summary = """FleetTrack Pro is a comprehensive, enterprise-grade fleet management and telematics SaaS platform designed to help transportation companies optimize their operations, ensure regulatory compliance, and improve safety across their entire fleet.

The platform provides real-time vehicle tracking, Electronic Logging Device (ELD) compliance management, driver performance monitoring, incident reporting, and advanced analytics - all within a secure, multi-tenant architecture that scales from small fleets to enterprise operations.

Key Business Objectives:
• Reduce operational costs through optimized routing and fuel management
• Ensure DOT/FMCSA compliance with automated ELD logging and reporting
• Improve fleet safety through real-time monitoring and incident management
• Provide actionable insights through comprehensive analytics and reporting
• Enable multi-organization support with role-based access control

Target Users:
• Fleet Managers - Oversee vehicle and driver operations
• Drivers - Log hours, view assignments, and report incidents
• Administrators - Manage organizations, users, and system settings
• Safety Officers - Monitor safety events and manage incidents"""

    doc.add_paragraph(exec_summary)
    doc.add_page_break()

    # ============ 2. FEATURE LIST ============
    doc.add_heading("2. Feature List", 1)

    # 2.1 Core Fleet Management
    doc.add_heading("2.1 Core Fleet Management", 2)

    features_table = doc.add_table(rows=1, cols=3)
    features_table.style = 'Table Grid'

    header_cells = features_table.rows[0].cells
    headers = ['Feature', 'Description', 'Capabilities']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '0066CC')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    fleet_features = [
        ("Dashboard", "Central command center for fleet operations", "• Real-time fleet statistics\n• Vehicle status overview\n• Alert notifications\n• Quick action cards"),
        ("Live Tracking", "Real-time GPS tracking of all vehicles", "• Interactive map view\n• Vehicle location markers\n• Telemetry data display\n• Geofence visualization"),
        ("Vehicle Management", "Complete vehicle lifecycle management", "• Add/Edit/Delete vehicles\n• VIN and registration tracking\n• Maintenance scheduling\n• Status management"),
        ("Driver Management", "Driver profiles and assignment management", "• Driver profiles with licenses\n• DOT medical certification\n• Vehicle assignments\n• Performance tracking"),
        ("Trip Tracking", "Comprehensive trip recording and analysis", "• Start/End locations\n• Distance and duration\n• Fuel consumption\n• Speed metrics"),
        ("Geofencing", "Virtual boundary management", "• Circle and polygon zones\n• Entry/Exit alerts\n• Custom alert rules\n• Zone visualization"),
    ]

    for feature, desc, caps in fleet_features:
        row = features_table.add_row()
        row.cells[0].text = feature
        row.cells[1].text = desc
        row.cells[2].text = caps
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # 2.2 ELD Compliance
    doc.add_heading("2.2 ELD Compliance (Electronic Logging Device)", 2)

    eld_table = doc.add_table(rows=1, cols=3)
    eld_table.style = 'Table Grid'

    header_cells = eld_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '0066CC')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    eld_features = [
        ("Hours of Service", "DOT-compliant HOS tracking", "• Off-duty logging\n• Sleeper berth tracking\n• Driving time\n• On-duty not driving"),
        ("Log Certification", "Daily log certification workflow", "• Driver self-certification\n• Compliance verification\n• Historical log access\n• Audit trail"),
        ("Violation Detection", "Automatic HOS violation alerts", "• Real-time monitoring\n• Violation notifications\n• Compliance reports\n• Remediation tracking"),
    ]

    for feature, desc, caps in eld_features:
        row = eld_table.add_row()
        row.cells[0].text = feature
        row.cells[1].text = desc
        row.cells[2].text = caps
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # 2.3 Safety & Analytics
    doc.add_heading("2.3 Safety & Analytics", 2)

    safety_table = doc.add_table(rows=1, cols=3)
    safety_table.style = 'Table Grid'

    header_cells = safety_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '0066CC')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    safety_features = [
        ("Incident Management", "Complete incident lifecycle handling", "• Incident reporting\n• Status tracking\n• Resolution workflow\n• Video attachment"),
        ("Safety Events", "Behavioral event monitoring", "• Harsh braking detection\n• Rapid acceleration\n• Speeding alerts\n• Distraction/Fatigue"),
        ("Alert System", "Comprehensive notification system", "• Geofence alerts\n• Maintenance reminders\n• ELD violations\n• Safety notifications"),
        ("Analytics Dashboard", "Data-driven fleet insights", "• Performance metrics\n• Fuel analysis\n• Driver rankings\n• Cost tracking"),
        ("Video Management", "Dashcam and incident footage", "• Video uploads\n• Incident association\n• Evidence management\n• Playback"),
    ]

    for feature, desc, caps in safety_features:
        row = safety_table.add_row()
        row.cells[0].text = feature
        row.cells[1].text = desc
        row.cells[2].text = caps
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # 2.4 Administration
    doc.add_heading("2.4 Administration & Security", 2)

    admin_table = doc.add_table(rows=1, cols=3)
    admin_table.style = 'Table Grid'

    header_cells = admin_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '0066CC')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    admin_features = [
        ("Multi-Tenancy", "Organization-based data isolation", "• Separate organizations\n• Subscription tiers\n• Custom branding\n• Data isolation"),
        ("User Management", "Role-based access control", "• 8 user roles\n• Granular permissions\n• User provisioning\n• Activity tracking"),
        ("Authentication", "Secure login system", "• Email/password auth\n• JWT sessions\n• 2FA support\n• Session management"),
        ("Audit Logging", "Comprehensive activity tracking", "• User action logs\n• System changes\n• Compliance reports\n• Data export"),
        ("Settings", "User preference management", "• Profile settings\n• Notification prefs\n• Password management\n• Theme options"),
    ]

    for feature, desc, caps in admin_features:
        row = admin_table.add_row()
        row.cells[0].text = feature
        row.cells[1].text = desc
        row.cells[2].text = caps
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ============ 3. TECHNOLOGY STACK ============
    doc.add_heading("3. Technology Stack", 1)

    # Frontend
    doc.add_heading("3.1 Frontend Technologies", 2)

    frontend_table = doc.add_table(rows=1, cols=3)
    frontend_table.style = 'Table Grid'

    header_cells = frontend_table.rows[0].cells
    tech_headers = ['Technology', 'Version', 'Purpose']
    for i, header in enumerate(tech_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2ECC71')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    frontend_tech = [
        ("Next.js", "16.0.10", "React framework with server-side rendering and API routes"),
        ("React", "19.2.1", "Component-based UI library"),
        ("Tailwind CSS", "4.0", "Utility-first CSS framework"),
        ("Leaflet", "1.9.4", "Interactive mapping library"),
        ("Recharts", "3.6.0", "Data visualization and charting"),
        ("Zustand", "5.0.9", "Lightweight state management"),
        ("TanStack Query", "5.90.12", "Data fetching and caching"),
        ("Lucide React", "0.561.0", "Icon library"),
    ]

    for tech, version, purpose in frontend_tech:
        row = frontend_table.add_row()
        row.cells[0].text = tech
        row.cells[1].text = version
        row.cells[2].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Backend
    doc.add_heading("3.2 Backend Technologies", 2)

    backend_table = doc.add_table(rows=1, cols=3)
    backend_table.style = 'Table Grid'

    header_cells = backend_table.rows[0].cells
    for i, header in enumerate(tech_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'F39C12')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    backend_tech = [
        ("Node.js", "Latest LTS", "JavaScript runtime environment"),
        ("NextAuth.js", "5.0.0-beta", "Authentication framework"),
        ("Prisma", "5.22.0", "Type-safe ORM for database access"),
        ("PostgreSQL", "Latest", "Primary relational database"),
        ("Zod", "4.2.0", "Schema validation library"),
        ("bcryptjs", "3.0.3", "Password hashing"),
        ("jsonwebtoken", "9.0.3", "JWT token handling"),
        ("Pino", "10.1.0", "High-performance logging"),
    ]

    for tech, version, purpose in backend_tech:
        row = backend_table.add_row()
        row.cells[0].text = tech
        row.cells[1].text = version
        row.cells[2].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # DevOps
    doc.add_heading("3.3 DevOps & Infrastructure", 2)

    devops_table = doc.add_table(rows=1, cols=3)
    devops_table.style = 'Table Grid'

    header_cells = devops_table.rows[0].cells
    for i, header in enumerate(tech_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '9B59B6')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    devops_tech = [
        ("Vercel", "Latest", "Serverless deployment platform"),
        ("Sentry", "10.32.1", "Error tracking and monitoring"),
        ("TypeScript", "5.x", "Type-safe JavaScript"),
        ("ESLint", "9.x", "Code linting and standards"),
        ("Vitest", "4.0.16", "Unit and integration testing"),
        ("Git", "Latest", "Version control"),
    ]

    for tech, version, purpose in devops_tech:
        row = devops_table.add_row()
        row.cells[0].text = tech
        row.cells[1].text = version
        row.cells[2].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ============ 4. DATA FLOW DIAGRAM ============
    doc.add_heading("4. Data Flow Diagram", 1)

    doc.add_paragraph("The following diagram illustrates how data flows through the FleetTrack Pro platform, from user interactions and IoT devices through the application layers to the database.")
    doc.add_paragraph()

    # Add the data flow diagram image
    try:
        doc.add_picture(os.path.join(OUTPUT_DIR, 'data_flow_diagram.png'), width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Data Flow Diagram - data_flow_diagram.png]")

    doc.add_paragraph()

    # Data flow description
    doc.add_heading("4.1 Data Flow Description", 2)

    flows = [
        ("User Authentication Flow", "Users submit credentials → NextAuth validates against database → Password verified with bcryptjs → JWT token created → Session established with 24-hour expiration → User redirected to dashboard"),
        ("Real-Time Tracking Flow", "Active vehicles queried → Location/telemetry data retrieved → Geofences fetched and parsed → Data transformed for map display → Frontend renders interactive map with markers"),
        ("Incident Reporting Flow", "User submits incident form → Incident record created → Alert automatically generated → Dashboard updated → Notifications dispatched"),
        ("Analytics Flow", "Time period specified → Multiple aggregations run in parallel → Trip stats, safety events, fuel data compiled → Driver performance calculated → Comprehensive analytics returned"),
    ]

    for title, desc in flows:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============ 5. SOFTWARE ARCHITECTURE ============
    doc.add_heading("5. Software Architecture", 1)

    doc.add_paragraph("FleetTrack Pro follows a layered architecture pattern with clear separation of concerns. The platform is built as a Multi-Tenant SaaS application with Role-Based Access Control (RBAC).")
    doc.add_paragraph()

    # Add architecture diagram
    try:
        doc.add_picture(os.path.join(OUTPUT_DIR, 'architecture_diagram.png'), width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph("[Architecture Diagram - architecture_diagram.png]")

    doc.add_paragraph()

    doc.add_heading("5.1 Architecture Layers", 2)

    layers = [
        ("Presentation Layer", "React-based UI components including Dashboard, Live Tracking, Fleet Management, ELD Compliance, Analytics, and Admin Panel. Built with Next.js pages and custom UI components."),
        ("Application Layer", "Next.js framework handling routing, middleware, API routes, authentication (NextAuth), and error handling with Sentry integration."),
        ("Business Logic Layer", "Core services including Fleet Management, Tracking, Trip Management, ELD Service, Safety Management, Analytics, and Geofencing."),
        ("Data Access Layer", "Prisma ORM providing type-safe database queries, migrations, and repository patterns for all 24 data models."),
        ("Infrastructure Layer", "PostgreSQL database, Vercel serverless hosting, security middleware (rate limiting, CORS), and external integrations (IoT devices, Maps)."),
    ]

    for layer, desc in layers:
        p = doc.add_paragraph()
        p.add_run(f"{layer}: ").bold = True
        p.add_run(desc)

    doc.add_heading("5.2 Key Architectural Patterns", 2)

    patterns = [
        "Multi-Tenancy: Organizations own all their data with complete isolation",
        "RBAC: 8 distinct user roles with granular permissions",
        "API-First: All data accessed through RESTful API endpoints",
        "Rate Limiting: Per-user and per-IP request throttling",
        "Request Correlation: UUID-based request tracing for debugging",
        "Structured Error Handling: Consistent error responses with logging",
        "Input Validation: Zod schemas validate all incoming data",
        "Session-Based Security: JWT tokens with configurable expiration",
    ]

    for pattern in patterns:
        doc.add_paragraph(f"• {pattern}")

    doc.add_page_break()

    # ============ 6. PLAN OF ACTION ============
    doc.add_heading("6. Plan of Action (Phase-wise Development)", 1)

    doc.add_paragraph("The development of FleetTrack Pro is organized into six strategic phases, each building upon the previous to deliver a complete, production-ready platform.")
    doc.add_paragraph()

    # Phase 1
    doc.add_heading("Phase 1: Foundation & Infrastructure", 2)

    phase1_table = doc.add_table(rows=1, cols=2)
    phase1_table.style = 'Table Grid'

    header_cells = phase1_table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Details"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '3498DB')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    phase1_items = [
        ("Project Setup", "Next.js 16 initialization, TypeScript configuration, ESLint/Prettier setup"),
        ("Database Design", "PostgreSQL schema design, Prisma ORM setup, migrations framework"),
        ("Authentication System", "NextAuth integration, JWT sessions, password hashing, RBAC foundation"),
        ("Core UI Framework", "Tailwind CSS setup, base components (Button, Input, Card, Table)"),
        ("API Infrastructure", "API route structure, error handling, logging (Pino), rate limiting"),
        ("Development Environment", "Local development setup, environment variables, Git workflow"),
    ]

    for item, details in phase1_items:
        row = phase1_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = details
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Phase 2
    doc.add_heading("Phase 2: Core Fleet Management", 2)

    phase2_table = doc.add_table(rows=1, cols=2)
    phase2_table.style = 'Table Grid'

    header_cells = phase2_table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Details"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '2ECC71')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    phase2_items = [
        ("Vehicle Management", "CRUD operations, vehicle profiles, status management, maintenance tracking"),
        ("Driver Management", "Driver profiles, license tracking, DOT medical dates, assignments"),
        ("Dashboard", "Statistics cards, fleet overview, quick actions, alert summary"),
        ("Organization Management", "Multi-tenant support, organization CRUD, subscription tiers"),
        ("User Management", "User CRUD, role assignment, permission enforcement"),
    ]

    for item, details in phase2_items:
        row = phase2_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = details
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Phase 3
    doc.add_heading("Phase 3: Real-Time Tracking & Telematics", 2)

    phase3_table = doc.add_table(rows=1, cols=2)
    phase3_table.style = 'Table Grid'

    header_cells = phase3_table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Details"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'F39C12')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    phase3_items = [
        ("Live Map Tracking", "Leaflet map integration, vehicle markers, real-time updates"),
        ("Telemetry Processing", "Engine data capture, sensor readings, diagnostic codes"),
        ("Geofencing System", "Zone creation (circle/polygon), entry/exit detection, alert triggers"),
        ("Trip Recording", "Start/end tracking, route logging, distance/duration calculation"),
        ("IoT Integration", "Device management, data ingestion API, telemetry storage"),
    ]

    for item, details in phase3_items:
        row = phase3_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = details
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Phase 4
    doc.add_heading("Phase 4: Compliance & Safety", 2)

    phase4_table = doc.add_table(rows=1, cols=2)
    phase4_table.style = 'Table Grid'

    header_cells = phase4_table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Details"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E74C3C')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    phase4_items = [
        ("ELD Logging", "Hours of Service tracking, duty status changes, log storage"),
        ("Log Certification", "Driver certification workflow, compliance verification"),
        ("Incident Management", "Incident reporting, status tracking, resolution workflow"),
        ("Safety Events", "Behavioral detection, event logging, driver scoring"),
        ("Alert System", "Multi-type alerts, notification delivery, acknowledgment tracking"),
    ]

    for item, details in phase4_items:
        row = phase4_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = details
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Phase 5
    doc.add_heading("Phase 5: Analytics & Reporting", 2)

    phase5_table = doc.add_table(rows=1, cols=2)
    phase5_table.style = 'Table Grid'

    header_cells = phase5_table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Details"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '9B59B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    phase5_items = [
        ("Analytics Dashboard", "Recharts integration, interactive visualizations, KPI displays"),
        ("Fleet Performance", "Trip analysis, fuel efficiency, utilization metrics"),
        ("Driver Scoring", "Performance ranking, safety scores, productivity metrics"),
        ("Cost Analysis", "Maintenance costs, fuel expenses, operational budgeting"),
        ("Custom Reports", "Configurable date ranges, data export, scheduled reports"),
    ]

    for item, details in phase5_items:
        row = phase5_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = details
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Phase 6
    doc.add_heading("Phase 6: Polish & Deployment", 2)

    phase6_table = doc.add_table(rows=1, cols=2)
    phase6_table.style = 'Table Grid'

    header_cells = phase6_table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Details"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '1ABC9C')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    phase6_items = [
        ("Testing & QA", "Unit tests (Vitest), integration tests, E2E testing, security audit"),
        ("Performance Optimization", "Database indexing, query optimization, caching strategies"),
        ("Security Hardening", "Penetration testing, vulnerability fixes, security headers"),
        ("Documentation", "API documentation, user guides, admin manuals"),
        ("Production Deployment", "Vercel configuration, environment setup, monitoring (Sentry)"),
        ("Launch Support", "Bug fixes, performance tuning, user feedback incorporation"),
    ]

    for item, details in phase6_items:
        row = phase6_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = details
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ============ 7. REQUIRED TEAM SETUP ============
    doc.add_heading("7. Required Team Setup", 1)

    doc.add_paragraph("The following team structure is recommended for successful delivery of the FleetTrack Pro platform:")
    doc.add_paragraph()

    team_table = doc.add_table(rows=1, cols=4)
    team_table.style = 'Table Grid'

    header_cells = team_table.rows[0].cells
    team_headers = ['Role', 'Count', 'Expertise Required', 'Responsibilities']
    for i, header in enumerate(team_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '34495E')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    team_roles = [
        ("Project Manager", "1", "Agile/Scrum, Fleet/Logistics domain", "Sprint planning, stakeholder communication, timeline management"),
        ("Tech Lead / Architect", "1", "Next.js, PostgreSQL, System Design", "Architecture decisions, code reviews, technical guidance"),
        ("Senior Full-Stack Developer", "2", "Next.js, React, Node.js, Prisma", "Core feature development, API design, complex integrations"),
        ("Full-Stack Developer", "2", "React, TypeScript, REST APIs", "Feature implementation, bug fixes, unit testing"),
        ("Frontend Developer", "1", "React, Tailwind CSS, Leaflet", "UI components, map integration, responsive design"),
        ("Backend Developer", "1", "Node.js, PostgreSQL, Authentication", "API development, database optimization, security"),
        ("QA Engineer", "1", "Vitest, E2E testing, API testing", "Test planning, automation, bug tracking"),
        ("DevOps Engineer", "1", "Vercel, CI/CD, Monitoring", "Deployment pipelines, infrastructure, monitoring"),
        ("UI/UX Designer", "1", "Figma, Design Systems, Prototyping", "User research, wireframes, visual design"),
    ]

    for role, count, expertise, resp in team_roles:
        row = team_table.add_row()
        row.cells[0].text = role
        row.cells[1].text = count
        row.cells[2].text = expertise
        row.cells[3].text = resp
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Team summary
    p = doc.add_paragraph()
    p.add_run("Total Team Size: ").bold = True
    p.add_run("11 members (can be scaled based on timeline requirements)")

    doc.add_paragraph()
    doc.add_heading("7.1 Optional/Specialized Roles", 2)

    optional_roles = [
        "IoT/Hardware Specialist - For ELD device integration and telemetry optimization",
        "Security Specialist - For compliance requirements (SOC 2, FMCSA)",
        "Database Administrator - For large-scale deployments with complex data requirements",
        "Technical Writer - For comprehensive documentation and API guides",
    ]

    for role in optional_roles:
        doc.add_paragraph(f"• {role}")

    doc.add_page_break()

    # ============ 8. PROJECT TIMELINE ============
    doc.add_heading("8. Project Timeline", 1)

    doc.add_paragraph("The following timeline represents a comprehensive development schedule for the FleetTrack Pro platform:")
    doc.add_paragraph()

    timeline_table = doc.add_table(rows=1, cols=4)
    timeline_table.style = 'Table Grid'

    header_cells = timeline_table.rows[0].cells
    timeline_headers = ['Phase', 'Duration', 'Key Milestones', 'Dependencies']
    for i, header in enumerate(timeline_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2C3E50')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    timeline_phases = [
        ("Phase 1: Foundation", "Weeks 1-4", "• Dev environment ready\n• Database schema complete\n• Auth system functional\n• Base UI components", "Team onboarding, requirements sign-off"),
        ("Phase 2: Core Fleet", "Weeks 5-10", "• Vehicle CRUD complete\n• Driver management done\n• Dashboard functional\n• Multi-tenancy working", "Phase 1 completion"),
        ("Phase 3: Tracking", "Weeks 11-16", "• Live map operational\n• Geofencing active\n• Trip logging working\n• Telemetry flowing", "Phase 2 completion, IoT setup"),
        ("Phase 4: Compliance", "Weeks 17-22", "• ELD logging active\n• Certification workflow\n• Incident management\n• Alert system live", "Phase 3 completion"),
        ("Phase 5: Analytics", "Weeks 23-26", "• Dashboards complete\n• Reports generating\n• Driver scores active\n• Cost tracking live", "Phase 4 completion"),
        ("Phase 6: Launch", "Weeks 27-30", "• All tests passing\n• Security audit passed\n• Documentation complete\n• Production deployed", "All phases complete"),
    ]

    for phase, duration, milestones, deps in timeline_phases:
        row = timeline_table.add_row()
        row.cells[0].text = phase
        row.cells[1].text = duration
        row.cells[2].text = milestones
        row.cells[3].text = deps
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Summary
    summary = doc.add_paragraph()
    summary.add_run("Total Estimated Duration: ").bold = True
    summary.add_run("30 weeks (approximately 7.5 months)")

    doc.add_paragraph()

    doc.add_heading("8.1 Timeline Considerations", 2)

    considerations = [
        "Timeline can be compressed with additional resources or parallel workstreams",
        "Buffer time is included for testing, bug fixes, and unexpected challenges",
        "Client feedback cycles may extend certain phases",
        "IoT device integration timeline depends on hardware availability",
        "Regulatory compliance requirements may require additional review cycles",
    ]

    for item in considerations:
        doc.add_paragraph(f"• {item}")

    doc.add_page_break()

    # ============ 9. API DOCUMENTATION ============
    doc.add_heading("9. API Documentation", 1)

    doc.add_paragraph("FleetTrack Pro exposes a comprehensive REST API with 30+ endpoints organized by functional domain. All endpoints require authentication except the health check endpoint.")
    doc.add_paragraph()

    # Authentication
    doc.add_heading("9.1 Authentication", 2)

    auth_table = doc.add_table(rows=1, cols=4)
    auth_table.style = 'Table Grid'

    header_cells = auth_table.rows[0].cells
    api_headers = ['Method', 'Endpoint', 'Description', 'Auth Required']
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '9B59B6')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    auth_endpoints = [
        ("POST", "/api/auth/[...nextauth]", "NextAuth handler (sign in/out)", "No"),
        ("GET", "/api/health", "Health check endpoint", "No"),
    ]

    for method, endpoint, desc, auth in auth_endpoints:
        row = auth_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Vehicles API
    doc.add_heading("9.2 Vehicles API", 2)

    vehicles_table = doc.add_table(rows=1, cols=4)
    vehicles_table.style = 'Table Grid'

    header_cells = vehicles_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '3498DB')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    vehicle_endpoints = [
        ("GET", "/api/vehicles", "List all vehicles (org-scoped)", "Yes"),
        ("POST", "/api/vehicles", "Create new vehicle", "Yes"),
        ("GET", "/api/vehicles/[id]", "Get vehicle with locations, telemetry", "Yes"),
        ("PATCH", "/api/vehicles/[id]", "Update vehicle status, odometer", "Yes"),
        ("DELETE", "/api/vehicles/[id]", "Delete vehicle (cascade)", "Yes"),
    ]

    for method, endpoint, desc, auth in vehicle_endpoints:
        row = vehicles_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Drivers API
    doc.add_heading("9.3 Drivers API", 2)

    drivers_table = doc.add_table(rows=1, cols=4)
    drivers_table.style = 'Table Grid'

    header_cells = drivers_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2ECC71')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    driver_endpoints = [
        ("GET", "/api/drivers", "List all drivers with stats", "Yes"),
        ("POST", "/api/drivers", "Create driver (creates user + driver)", "Yes"),
        ("GET", "/api/drivers/[id]", "Get driver with trips and ELD logs", "Yes"),
        ("PATCH", "/api/drivers/[id]", "Update driver status, license info", "Yes"),
        ("DELETE", "/api/drivers/[id]", "Delete driver and related records", "Yes"),
    ]

    for method, endpoint, desc, auth in driver_endpoints:
        row = drivers_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ELD API
    doc.add_heading("9.4 ELD Compliance API", 2)

    eld_api_table = doc.add_table(rows=1, cols=4)
    eld_api_table.style = 'Table Grid'

    header_cells = eld_api_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'E74C3C')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    eld_endpoints = [
        ("GET", "/api/eld", "List ELD logs (date-filtered)", "Yes"),
        ("POST", "/api/eld", "Create ELD log entry", "Yes"),
        ("POST", "/api/eld/certify", "Certify daily ELD logs", "Yes"),
        ("GET", "/api/eld/status", "Get ELD status for driver", "Yes"),
    ]

    for method, endpoint, desc, auth in eld_endpoints:
        row = eld_api_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Additional APIs
    doc.add_heading("9.5 Additional API Endpoints", 2)

    additional_table = doc.add_table(rows=1, cols=4)
    additional_table.style = 'Table Grid'

    header_cells = additional_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'F39C12')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    additional_endpoints = [
        ("GET", "/api/tracking", "Get active vehicles with geofences", "Yes"),
        ("GET/POST", "/api/trips", "List/Create trips", "Yes"),
        ("GET/POST/PATCH/DELETE", "/api/incidents", "Incident management", "Yes"),
        ("GET/POST/PATCH/DELETE", "/api/geofences", "Geofence management", "Yes"),
        ("GET/PATCH", "/api/alerts", "Alert management", "Yes"),
        ("GET", "/api/analytics", "Fleet analytics", "Yes"),
        ("POST/PATCH/DELETE", "/api/admin/organizations", "Organization management", "Admin"),
        ("POST/PATCH/DELETE", "/api/admin/users", "User management", "Admin"),
        ("PATCH", "/api/settings/*", "User settings (profile, password, 2FA)", "Yes"),
    ]

    for method, endpoint, desc, auth in additional_endpoints:
        row = additional_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Response formats
    doc.add_heading("9.6 Response Formats", 2)

    p = doc.add_paragraph()
    p.add_run("Success Response:").bold = True

    success_code = """
{
  "id": "string",
  "licensePlate": "string",
  "make": "string",
  "status": "active" | "inactive" | "maintenance" | "out_of_service",
  // ... additional fields
}"""

    code_para = doc.add_paragraph(success_code)
    code_para.style = 'No Spacing'

    p = doc.add_paragraph()
    p.add_run("Error Response:").bold = True

    error_code = """
{
  "error": "string",
  "details": [...],  // optional validation errors
  "message": "string"  // optional
}"""

    code_para = doc.add_paragraph(error_code)
    code_para.style = 'No Spacing'

    doc.add_page_break()

    # ============ 10. DATABASE SCHEMA ============
    doc.add_heading("10. Database Schema", 1)

    doc.add_paragraph("FleetTrack Pro uses PostgreSQL with Prisma ORM, featuring 24 interconnected data models organized into logical domains.")
    doc.add_paragraph()

    # Organization & Users
    doc.add_heading("10.1 Organization & Users", 2)

    org_table = doc.add_table(rows=1, cols=4)
    org_table.style = 'Table Grid'

    header_cells = org_table.rows[0].cells
    schema_headers = ['Table', 'Key Fields', 'Relationships', 'Purpose']
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '8E44AD')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    org_models = [
        ("Organization", "id, name, slug, subscriptionTier, status", "Has many: Users, Vehicles, Drivers, Geofences", "Multi-tenant container"),
        ("User", "id, email, password, role, preferences, twoFactorEnabled", "Belongs to: Organization; Has one: Driver", "User accounts and auth"),
        ("Session", "id, token, userId, expiresAt", "Belongs to: User", "JWT session tracking"),
    ]

    for table, fields, rels, purpose in org_models:
        row = org_table.add_row()
        row.cells[0].text = table
        row.cells[1].text = fields
        row.cells[2].text = rels
        row.cells[3].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Fleet Management
    doc.add_heading("10.2 Fleet Management", 2)

    fleet_table = doc.add_table(rows=1, cols=4)
    fleet_table.style = 'Table Grid'

    header_cells = fleet_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '3498DB')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    fleet_models = [
        ("Vehicle", "id, vin, licensePlate, make, model, year, status, odometer", "Belongs to: Organization; Has many: Locations, Telemetry, Trips", "Vehicle profiles"),
        ("Driver", "id, licenseNumber, licenseState, licenseExpiry, dotMedicalExpiry, status", "Belongs to: User, Organization; Has many: Trips, ELDLogs", "Driver profiles"),
        ("VehicleAssignment", "id, vehicleId, driverId, startDate, endDate", "Belongs to: Vehicle, Driver", "Assignment history"),
        ("IOTDevice", "id, serialNumber, type, status, firmwareVersion", "Belongs to: Vehicle", "GPS/ELD devices"),
    ]

    for table, fields, rels, purpose in fleet_models:
        row = fleet_table.add_row()
        row.cells[0].text = table
        row.cells[1].text = fields
        row.cells[2].text = rels
        row.cells[3].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Tracking & Telemetry
    doc.add_heading("10.3 Tracking & Telemetry", 2)

    tracking_table = doc.add_table(rows=1, cols=4)
    tracking_table.style = 'Table Grid'

    header_cells = tracking_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2ECC71')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tracking_models = [
        ("VehicleLocation", "id, vehicleId, latitude, longitude, speed, heading, timestamp", "Belongs to: Vehicle; Index: (vehicleId, timestamp)", "GPS coordinates"),
        ("VehicleTelemetry", "id, vehicleId, engineRpm, fuelLevel, coolantTemp, oilPressure", "Belongs to: Vehicle; Index: (vehicleId, timestamp)", "Engine metrics"),
        ("VehicleDiagnostic", "id, vehicleId, code, description, severity, status", "Belongs to: Vehicle", "OBD-II codes"),
        ("Trip", "id, vehicleId, driverId, startLat, startLng, endLat, endLng, distance, duration", "Belongs to: Vehicle, Driver", "Trip records"),
        ("Route", "id, name, waypoints, estimatedDuration, estimatedDistance", "Belongs to: Organization", "Pre-defined routes"),
    ]

    for table, fields, rels, purpose in tracking_models:
        row = tracking_table.add_row()
        row.cells[0].text = table
        row.cells[1].text = fields
        row.cells[2].text = rels
        row.cells[3].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Compliance & Safety
    doc.add_heading("10.4 Compliance & Safety", 2)

    compliance_table = doc.add_table(rows=1, cols=4)
    compliance_table.style = 'Table Grid'

    header_cells = compliance_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'E74C3C')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    compliance_models = [
        ("ELDLog", "id, driverId, date, status, startTime, endTime, duration, certified", "Belongs to: Driver; Index: (driverId, date)", "Hours of Service"),
        ("SafetyEvent", "id, driverId, vehicleId, type, severity, location, timestamp", "Belongs to: Driver, Vehicle; Index: (driverId, timestamp)", "Behavioral events"),
        ("Incident", "id, driverId, vehicleId, type, severity, status, location, description", "Belongs to: Driver, Vehicle, Organization", "Incident reports"),
        ("Video", "id, incidentId, vehicleId, url, duration, type", "Belongs to: Incident, Vehicle", "Dashcam footage"),
        ("Alert", "id, type, severity, message, isRead, acknowledgedAt", "Belongs to: Organization", "System notifications"),
    ]

    for table, fields, rels, purpose in compliance_models:
        row = compliance_table.add_row()
        row.cells[0].text = table
        row.cells[1].text = fields
        row.cells[2].text = rels
        row.cells[3].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Operations
    doc.add_heading("10.5 Operations & Maintenance", 2)

    ops_table = doc.add_table(rows=1, cols=4)
    ops_table.style = 'Table Grid'

    header_cells = ops_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'F39C12')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    ops_models = [
        ("MaintenanceRecord", "id, vehicleId, type, description, cost, laborHours, status", "Belongs to: Vehicle", "Service records"),
        ("FuelRecord", "id, vehicleId, driverId, gallons, pricePerGallon, totalCost, odometer", "Belongs to: Vehicle, Driver", "Fuel purchases"),
        ("Geofence", "id, name, type, coordinates, alertOnEntry, alertOnExit", "Belongs to: Organization", "Virtual boundaries"),
        ("AuditLog", "id, userId, action, resource, resourceId, changes, timestamp", "Belongs to: User", "Activity tracking"),
    ]

    for table, fields, rels, purpose in ops_models:
        row = ops_table.add_row()
        row.cells[0].text = table
        row.cells[1].text = fields
        row.cells[2].text = rels
        row.cells[3].text = purpose
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Database indexes
    doc.add_heading("10.6 Performance Indexes", 2)

    indexes = [
        "VehicleLocation: (vehicleId, timestamp) - Optimized for latest location queries",
        "VehicleTelemetry: (vehicleId, timestamp) - Optimized for latest telemetry",
        "ELDLog: (driverId, date) - Optimized for daily compliance reports",
        "SafetyEvent: (driverId, timestamp) - Optimized for driver safety scoring",
        "Trip: (vehicleId, startTime) - Optimized for trip history queries",
    ]

    for index in indexes:
        doc.add_paragraph(f"• {index}")

    doc.add_page_break()

    # ============ 11. APPENDIX ============
    doc.add_heading("11. Appendix", 1)

    # User Roles
    doc.add_heading("11.1 User Roles & Permissions", 2)

    roles_table = doc.add_table(rows=1, cols=3)
    roles_table.style = 'Table Grid'

    header_cells = roles_table.rows[0].cells
    role_headers = ['Role', 'Level', 'Permissions']
    for i, header in enumerate(role_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '34495E')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    roles = [
        ("saas_admin", "Platform", "Full system access, manage all organizations"),
        ("saas_subscriber_manager", "Platform", "Manage subscriptions and billing"),
        ("saas_support", "Platform", "Customer support, read-only access"),
        ("saas_developer", "Platform", "API access, debugging tools"),
        ("company_admin", "Organization", "Full org access, manage users and settings"),
        ("fleet_manager", "Organization", "Manage vehicles, drivers, and operations"),
        ("driver", "Organization", "View own trips, ELD logs, report incidents"),
        ("company_support", "Organization", "Support staff, limited read access"),
    ]

    for role, level, perms in roles:
        row = roles_table.add_row()
        row.cells[0].text = role
        row.cells[1].text = level
        row.cells[2].text = perms
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Environment Variables
    doc.add_heading("11.2 Environment Configuration", 2)

    env_table = doc.add_table(rows=1, cols=3)
    env_table.style = 'Table Grid'

    header_cells = env_table.rows[0].cells
    env_headers = ['Variable', 'Required', 'Description']
    for i, header in enumerate(env_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '7F8C8D')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    env_vars = [
        ("DATABASE_URL", "Yes", "PostgreSQL connection string"),
        ("NEXTAUTH_URL", "Yes", "Application base URL"),
        ("NEXTAUTH_SECRET", "Yes", "JWT signing secret (32+ chars)"),
        ("AUTH_URL", "Yes", "Authentication callback URL"),
        ("LOG_LEVEL", "No", "Logging level (info, debug, error)"),
        ("RATE_LIMIT_MAX", "No", "Max requests per window (default: 100)"),
        ("RATE_LIMIT_WINDOW_MS", "No", "Rate limit window in ms (default: 60000)"),
    ]

    for var, req, desc in env_vars:
        row = env_table.add_row()
        row.cells[0].text = var
        row.cells[1].text = req
        row.cells[2].text = desc
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Security Features
    doc.add_heading("11.3 Security Features", 2)

    security = [
        "Password hashing with bcryptjs (12 salt rounds)",
        "JWT-based sessions with 24-hour expiration",
        "Rate limiting: 100 requests/60s (API), 5 requests/60s (auth)",
        "CORS validation on all API routes",
        "Security headers (X-Content-Type-Options, X-Frame-Options, X-XSS-Protection)",
        "Request ID correlation for audit trails",
        "Sensitive field redaction in logs",
        "Input validation with Zod schemas",
        "Two-factor authentication support",
    ]

    for item in security:
        doc.add_paragraph(f"• {item}")

    doc.add_paragraph()

    # Contact
    doc.add_heading("11.4 Document Information", 2)

    doc.add_paragraph("This Statement of Work provides a comprehensive overview of the FleetTrack Pro platform. For additional information, technical discussions, or project inquiries, please contact the project team.")

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("— End of Document —")
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.font.size = Pt(10)

    # Save document
    doc_path = os.path.join(OUTPUT_DIR, 'FleetTrack_Pro_Statement_of_Work.docx')
    doc.save(doc_path)
    print(f"Created: {doc_path}")
    return doc_path

def main():
    print("Generating FleetTrack Pro Statement of Work...")
    print("-" * 50)

    # Generate diagrams first
    print("\n1. Creating Data Flow Diagram...")
    create_data_flow_diagram()

    print("\n2. Creating Architecture Diagram...")
    create_architecture_diagram()

    print("\n3. Creating Statement of Work Document...")
    doc_path = create_sow_document()

    print("\n" + "=" * 50)
    print("Generation Complete!")
    print("=" * 50)
    print(f"\nFiles created in: {OUTPUT_DIR}")
    print("  - data_flow_diagram.png")
    print("  - architecture_diagram.png")
    print("  - FleetTrack_Pro_Statement_of_Work.docx")

if __name__ == "__main__":
    main()
